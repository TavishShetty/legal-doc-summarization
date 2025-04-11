from fastapi import FastAPI, UploadFile, File
from fastapi.middleware.cors import CORSMiddleware
import docx
import PyPDF2
import io
import base64

app = FastAPI()

app.add_middleware(
    CORSMiddleware,
    allow_origins=["http://localhost:8080"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

@app.post("/process")
async def process_document(file: UploadFile = File(...), summarize: bool = True, anonymize: bool = True):
    # Read file content
    content = ""
    if file.filename.endswith(".docx"):
        doc = docx.Document(file.file)
        content = " ".join([para.text for para in doc.paragraphs])
    elif file.filename.endswith(".pdf"):
        pdf = PyPDF2.PdfReader(file.file)
        content = " ".join([pdf.pages[i].extract_text() for i in range(len(pdf.pages))])
    elif file.filename.endswith(".txt"):
        content = (await file.read()).decode("utf-8")

    # Mock processing
    result = {}
    if summarize:
        result["summary"] = f"Summary of {file.filename}: {content[:50]}..."
    if anonymize:
        anonymized_text = content.replace("person", "[ANONYMIZED]").replace("name", "[NAME]")
        result["anonymized"] = anonymized_text
        # Return mocked file in base64
        if file.filename.endswith(".docx"):
            new_doc = docx.Document()
            new_doc.add_paragraph(anonymized_text)
            buffer = io.BytesIO()
            new_doc.save(buffer)
            result["anonymized_file"] = base64.b64encode(buffer.getvalue()).decode("utf-8")
        elif file.filename.endswith(".pdf"):
            # Mock as text (PDF generation skipped for simplicity)
            result["anonymized_file"] = base64.b64encode(anonymized_text.encode("utf-8")).decode("utf-8")
        else:
            result["anonymized_file"] = base64.b64encode(anonymized_text.encode("utf-8")).decode("utf-8")

    return result

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)